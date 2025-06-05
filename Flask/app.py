from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from transformers import pipeline, AutoTokenizer
import torch
import os
from werkzeug.utils import secure_filename
import PyPDF2
import docx
import re
import logging
import uuid
from concurrent.futures import ThreadPoolExecutor
import time
import shutil
import traceback
import random
from functools import lru_cache, wraps
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings

import requests
from datetime import datetime, timedelta
from werkzeug.exceptions import HTTPException
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Initialize Flask app
app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Model configuration for summarization
MODEL_NAME = "Ruthwik/LExiMinD_legal_t5_summarizer"
CHUNK_SIZE = 1024  # Size for parallel processing
DEFAULT_MAX_LENGTH = 300  # Balanced default length
DEFAULT_MIN_LENGTH = 100

# Configuration for file handling
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
PREPROCESSED_FOLDER = os.path.join(BASE_DIR, 'preprocessed')
PROCESSED_FOLDER = os.path.join(BASE_DIR, 'processed')

# Create directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PREPROCESSED_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PREPROCESSED_FOLDER'] = PREPROCESSED_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB file size limit

# Legal QA configuration
LEGAL_TERMS = [
    # Court and judicial terms
    "court", "judgment", "judge", "bench", "hon'ble", "honorable", "justice",
    "chief justice", "justice", "j.", "cj", "division bench", "full bench",
    "single bench", "coram", "jurisdiction", "original jurisdiction",
    "appellate jurisdiction", "writ jurisdiction", "revision", "review",
    "curative petition", "special leave petition", "slp", "civil appeal",
    "criminal appeal", "letters patent appeal", "lpa", "review petition",

    # Judgment document structure
    "case no.", "in the matter of", "versus", "vs", "v.", "petitioner",
    "respondent", "appellant", "appellee", "complainant", "accused",
    "defendant", "plaintiff", "applicant", "opposite party", "op",
    "intervenor", "amicus curiae", "next friend", "pro forma respondent",

    # Legal document sections
    "headnote", "citation", "facts", "issues", "arguments", "submissions",
    "contentions", "pleadings", "evidence", "exhibits", "affidavit",
    "deposition", "testimony", "witness", "examination", "cross-examination",
    "reexamination", "documents", "annexures", "schedules", "appendices",
    "preamble", "recitals", "operative portion", "ratio decidendi",
    "obiter dicta", "holding", "findings", "conclusions", "decision",
    "order", "decree", "final order", "interim order", "injunction",
    "stay", "bail", "remand", "custody", "parole", "probation",

    # Legal procedures
    "filing", "institution", "commencement", "cause of action",
    "limitation", "prescription", "res judicata", "lis pendens",
    "sub judice", "stare decisis", "precedent", "binding precedent",
    "persuasive precedent", "distinguished", "overruled", "reversed",
    "affirmed", "modified", "remanded", "disposed", "dismissed",
    "allowed", "partly allowed", "quashed", "set aside", "annulled",
    "struck down", "upheld", "sustained", "vacated", "withdrawn",

    # Constitutional law
    "constitution", "constitutional", "unconstitutional", "ultra vires",
    "intra vires", "basic structure", "fundamental rights", "directive principles",
    "fundamental duties", "writ", "habeas corpus", "mandamus", "prohibition",
    "certiorari", "quo warranto", "article 14", "article 19", "article 21",
    "article 32", "article 226", "article 136", "article 142", "article 144",
    "separation of powers", "judicial review", "rule of law", "due process",
    "equal protection", "reasonable restriction", "public interest",
    "doctrine of eclipse", "doctrine of severability", "colourable legislation",

    # Civil law
    "civil procedure code", "cpc", "order", "rule", "section", "appeal",
    "revision", "review", "execution", "decree", "judgment", "plaint",
    "written statement", "counter claim", "set off", "interlocutory",
    "interim relief", "injunction", "temporary injunction", "permanent injunction",
    "specific performance", "declaration", "damages", "compensation",
    "mesne profits", "restitution", "receiver", "commission", "discovery",
    "interrogatories", "admission", "denial", "affidavit", "ex parte",
    "ex parte decree", "ex parte order", "setting aside ex parte",

    # Criminal law
    "criminal procedure code", "crpc", "indian penal code", "ipc",
    "section 302", "section 304", "section 307", "section 376", "section 420",
    "section 498a", "bailable", "non-bailable", "cognizable", "non-cognizable",
    "fir", "charge sheet", "charges", "framing of charges", "discharge",
    "acquittal", "conviction", "sentence", "death sentence", "life imprisonment",
    "fine", "compensation", "probation", "parole", "remission", "commutation",
    "suspension", "anticipatory bail", "regular bail", "default bail",
    "custodial interrogation", "police custody", "judicial custody",
    "remand", "discharge", "compounding", "quashing", "stay", "suspension",

    # Evidence law
    "indian evidence act", "evidence", "proof", "burden of proof",
    "onus of proof", "standard of proof", "presumption", "rebuttable",
    "irrebuttable", "documentary evidence", "oral evidence", "expert evidence",
    "circumstantial evidence", "direct evidence", "hearsay", "confession",
    "admission", "dying declaration", "hostile witness", "leading question",
    "cross examination", "reexamination", "affidavit", "exhibit", "marking",
    "identification", "proof of document", "secondary evidence", "primary evidence",

    # Contract law
    "indian contract act", "contract", "agreement", "offer", "acceptance",
    "consideration", "competent parties", "free consent", "coercion",
    "undue influence", "fraud", "misrepresentation", "mistake", "void",
    "voidable", "unenforceable", "quasi contract", "specific performance",
    "damages", "compensation", "liquidated damages", "penalty", "rescission",
    "rectification", "restitution", "quantum meruit", "breach", "anticipatory",
    "actual breach", "remedies", "injunction", "declaration",

    # Property law
    "transfer of property act", "topa", "sale", "mortgage", "lease",
    "gift", "exchange", "license", "easement", "adverse possession",
    "prescription", "title", "ownership", "possession", "constructive possession",
    "joint possession", "co-ownership", "coparcenary", "partition", "will",
    "testament", "codicil", "probate", "letters of administration",
    "succession", "inheritance", "heir", "legatee", "devisee", "bequest",
    "gift", "settlement", "trust", "beneficiary", "trustee", "endowment",

    # Company law
    "companies act", "memorandum", "articles", "incorporation", "registration",
    "director", "managing director", "whole-time director", "independent director",
    "nominee director", "board of directors", "general meeting", "agm", "egm",
    "resolution", "ordinary resolution", "special resolution", "shareholder",
    "member", "share", "equity share", "preference share", "debenture",
    "charge", "mortgage", "lien", "floating charge", "fixed charge", "winding up",
    "voluntary winding up", "compulsory winding up", "liquidation", "official liquidator",
    "insolvency", "bankruptcy", "resolution professional", "liquidator",
    "winding up petition", "oppression and mismanagement", "nclt", "nclat",

    # Intellectual property
    "patent", "copyright", "trademark", "design", "geographical indication",
    "infringement", "passing off", "counterfeiting", "piracy", "plagiarism",
    "assignment", "license", "compulsory license", "royalty", "damages",
    "injunction", "anticipatory injunction", "permanent injunction",
    "account of profits", "seizure", "destruction", "rectification",

    # Labor and industrial law
    "industrial disputes act", "ida", "workman", "employer", "employee",
    "industrial dispute", "strike", "lockout", "layoff", "retrenchment",
    "closure", "transfer", "closure", "compensation", "gratuity", "bonus",
    "provident fund", "esic", "epf", "minimum wages", "equal remuneration",
    "sexual harassment", "disciplinary proceedings", "domestic enquiry",
    "punishment", "dismissal", "termination", "reinstatement", "back wages",
    "conciliation", "arbitration", "adjudication", "labour court",
    "industrial tribunal", "national tribunal", "collective bargaining",
    "settlement", "award", "implementation", "enforcement",

    # Tax law
    "income tax act", "gst", "vat", "customs", "excise", "service tax",
    "assessment", "reassessment", "scrutiny", "regular assessment",
    "best judgment assessment", "appeal", "revision", "rectification",
    "advance ruling", "settlement commission", "tax evasion", "tax avoidance",
    "penalty", "prosecution", "recovery", "attachment", "garnishee",
    "stay", "refund", "tribunal", "high court", "supreme court",

    # Arbitration
    "arbitration", "conciliation", "mediation", "award", "enforcement",
    "setting aside", "arbitrator", "umpire", "arbitral tribunal",
    "arbitration agreement", "seat", "venue", "jurisdiction", "competence",
    "competence-competence", "interim measures", "emergency arbitrator",
    "final award", "partial award", "interest", "costs", "challenge",
    "neutrality", "impartiality", "independence", "disclosure",

    # International law
    "treaty", "convention", "protocol", "united nations", "general assembly",
    "security council", "international court of justice", "pcij", "icj",
    "arbitration", "mediation", "good offices", "diplomatic protection",
    "state responsibility", "immunity", "sovereign immunity", "jurisdiction",
    "extradition", "mutual legal assistance", "human rights", "refugee",
    "asylum", "extra-territorial", "exhaustion of local remedies",

    # Legal maxims
    "actus reus", "mens rea", "audi alteram partem", "nemo judex in causa sua",
    "res ipsa loquitur", "uberrima fides", "caveat emptor", "stare decisis",
    "obiter dicta", "ratio decidendi", "ignorantia juris non excusat",
    "de minimis non curat lex", "expressio unius est exclusio alterius",
    "ejusdem generis", "noscitur a sociis", "pari materia", "in pari delicto",
    "volenti non fit injuria", "damnum sine injuria", "injuria sine damno",
    "qui facit per alium facit per se", "respondeat superior", "volenti non fit injuria",

    # Case-specific terms (from your judgment)
    "appeal no. 790 of 1957", "civil misc writ no. 280 of 1950",
    "u.p. industrial disputes act", "xxviii of 1947", "court of inquiry",
    "allahabad high court", "mudholkar j.", "bhargava j.", "sapru j.",
    "state of uttar pradesh", "indian sugar millers association",
    "indian national sugar mills workers federation", "sugar factories",
    "bonus payment", "retaining allowance", "seasonal workmen", "clerical staff",
    "industrial dispute", "strike notice", "court of inquiry", "gazette notification",
    "writ petition", "article 226", "article 133", "full bench", "constitutional validity",
    "ultra vires", "discrimination", "arbitrary", "public interest", "emergency",
    "prospective", "retrospective", "minimum wages act", "collective bargaining",
    "terms of employment", "conditions of employment", "mandamus", "certificate",
    "special leave petition", "constitution bench",

    # Terms related to user questions and queries
    "explain", "clarify", "interpret", "define", "what does", "meaning of",
    "how to", "procedure for", "requirements for", "eligibility for",
    "criteria for", "difference between", "similarities between",
    "compare", "contrast", "examples of", "types of", "categories of",
    "applicability of", "scope of", "limitations of", "exceptions to",
    "validity of", "enforceability of", "consequences of", "penalty for",
    "remedy for", "solution for", "process for", "steps to", "guide to",
    "analysis of", "breakdown of", "summary of", "overview of",
    "key points", "main arguments", "legal basis", "grounds for",
    "justification for", "rationale behind", "purpose of", "intent behind",
    "objective of", "effect of", "impact of", "implications of",
    "significance of", "importance of", "relevance of", "connection between",
    "relationship between", "correlation between", "cause of", "effect of",
    "reason for", "basis for", "foundation of", "principle behind",
    "doctrine of", "theory of", "concept of", "aspects of", "elements of",
    "components of", "factors in", "considerations for", "requirements of",
    "conditions for", "terms of", "provisions of", "clauses in",
    "sections in", "articles in", "rules in", "regulations in",
    "guidelines for", "standards for", "benchmarks for", "precedents for",
    "case law on", "jurisprudence on", "legal opinion on", "view on",
    "position on", "stance on", "interpretation of", "construction of",
    "reading of", "understanding of", "comprehension of", "application of",
    "implementation of", "execution of", "enforcement of", "compliance with",
    "adherence to", "obligation to", "right to", "entitlement to",
    "privilege of", "immunity from", "exception to", "exemption from",
    "derogation from", "deviation from", "variation of", "modification of",
    "amendment to", "revision of", "update to", "change in",
    "development in", "trend in", "pattern in", "practice of",
    "custom of", "usage of", "tradition of", "convention of",
    "norm of", "standard of", "measure of", "test for",
    "criteria for", "benchmark for", "yardstick for", "indicator of",
    "evidence of", "proof of", "verification of", "confirmation of",
    "validation of", "authentication of", "certification of", "approval of",
    "authorization of", "sanction of", "ratification of", "endorsement of",
    "support for", "opposition to", "objection to", "challenge to",
    "appeal against", "review of", "reconsideration of", "revision of",
    "reformation of", "rectification of", "correction of", "amendment of",
    "modification of", "alteration of", "change to", "adjustment to",
    "adaptation of", "transformation of", "conversion of", "translation of",
    "paraphrase of", "summary of", "abstract of", "synopsis of",
    "outline of", "overview of", "introduction to", "background of",
    "context of", "framework of", "structure of", "organization of",
    "hierarchy of", "classification of", "categorization of", "typology of",
    "taxonomy of", "nomenclature of", "terminology of", "vocabulary of",
    "glossary of", "dictionary of", "lexicon of", "thesaurus of",
    "encyclopedia of", "compendium of", "digest of", "manual of",
    "handbook of", "guidebook of", "textbook of", "treatise on",
    "monograph on", "dissertation on", "thesis on", "paper on",
    "article on", "essay on", "commentary on", "annotation of",
    "exegesis of", "hermeneutics of", "interpretation of", "construction of",
    "reading of", "analysis of", "examination of", "investigation of",
    "inquiry into", "research on", "study of", "survey of",
    "report on", "finding of", "conclusion of", "recommendation of",
    "suggestion for", "proposal for", "plan for", "strategy for",
    "approach to", "method for", "technique for", "procedure for",
    "process for", "system for", "framework for", "model for",
    "paradigm for", "template for", "prototype for", "example of",
    "instance of", "case of", "illustration of", "demonstration of",
    "exposition of", "explanation of", "clarification of", "elucidation of",
    "simplification of", "breakdown of", "deconstruction of", "reconstruction of",
    "synthesis of", "integration of", "unification of", "harmonization of",
    "reconciliation of", "alignment of", "coordination of", "orchestration of",
    "management of", "administration of", "governance of", "regulation of",
    "control of", "supervision of", "oversight of", "monitoring of",
    "evaluation of", "assessment of", "appraisal of", "review of",
    "audit of", "inspection of", "scrutiny of", "examination of",
    "verification of", "validation of", "authentication of", "certification of",
    "accreditation of", "licensing of", "authorization of", "approval of",
    "sanction of", "endorsement of", "ratification of", "confirmation of",
    "affirmation of", "declaration of", "pronouncement of", "announcement of",
    "publication of", "dissemination of", "distribution of", "circulation of",
    "promulgation of", "enactment of", "legislation of", "regulation of",
    "ordinance of", "decree of", "edict of", "proclamation of",
    "notification of", "directive of", "instruction of", "order of",
    "command of", "injunction of", "mandate of", "requirement of",
    "obligation of", "duty of", "responsibility of", "accountability of",
    "liability of", "culpability of", "blameworthiness of", "fault of",
    "negligence of", "recklessness of", "intent of", "purpose of",
    "motive of", "reason for", "cause of", "origin of",
    "source of", "basis of", "foundation of", "ground of",
    "justification for", "rationale for", "explanation for", "defense of",
    "excuse for", "pretext for", "alibi for", "vindication of",
    "exoneration of", "absolution of", "acquittal of", "discharge of",
    "release from", "liberation from", "emancipation from", "freedom from",
    "exemption from", "exception to", "derogation from", "deviation from",
    "departure from", "variation from", "modification of", "alteration of",
    "change to", "adjustment to", "adaptation of", "transformation of",
    "conversion of", "translation of", "interpretation of", "construction of",
    "reading of", "understanding of", "comprehension of", "appreciation of",
    "recognition of", "acknowledgment of", "admission of", "concession of",
    "confession of", "disclosure of", "revelation of", "exposure of",
    "discovery of", "finding of", "determination of", "resolution of",
    "decision on", "judgment on", "ruling on", "verdict on",
    "sentence on", "order on", "decree on", "pronouncement on",
    "declaration on", "announcement on", "publication of", "issuance of",
    "delivery of", "service of", "filing of", "submission of",
    "presentation of", "tender of", "offer of", "proposal of",
    "suggestion of", "recommendation of", "advice on", "counsel on",
    "guidance on", "direction on", "instruction on", "command on",
    "order on", "injunction on", "mandate on", "requirement on",
    "demand for", "request for", "petition for", "application for",
    "appeal for", "plea for", "prayer for", "suit for",
    "action for", "case for", "matter of", "issue of",
    "question of", "point of", "aspect of", "element of",
    "factor in", "component of", "ingredient of", "feature of",
    "characteristic of", "attribute of", "quality of", "property of",
    "trait of", "mark of", "sign of", "indication of",
    "evidence of", "proof of", "verification of", "confirmation of",
    "validation of", "authentication of", "certification of", "attestation of",
    "witnessing of", "observation of", "perception of", "view of",
    "opinion of", "belief of", "conviction of", "position of",
    "stance of", "attitude toward", "approach to", "method for",
    "technique for", "procedure for", "process for", "system for",
    "framework for", "model for", "paradigm for", "template for",
    "prototype for", "example of", "instance of", "case of",
    "illustration of", "demonstration of", "exposition of", "explanation of",
    "clarification of", "elucidation of", "simplification of", "breakdown of",
    "deconstruction of", "reconstruction of", "synthesis of", "integration of",
    "unification of", "harmonization of", "reconciliation of", "alignment of",
    "coordination of", "orchestration of", "management of", "administration of",
    "governance of", "regulation of", "control of", "supervision of",
    "oversight of", "monitoring of", "evaluation of", "assessment of",
    "appraisal of", "review of", "audit of", "inspection of",
    "scrutiny of", "examination of", "verification of", "validation of",
    "authentication of", "certification of", "accreditation of", "licensing of",
    "authorization of", "approval of", "sanction of", "endorsement of",
    "ratification of", "confirmation of", "affirmation of", "declaration of",
    "pronouncement of", "announcement of", "publication of", "dissemination of",
    "distribution of", "circulation of", "promulgation of", "enactment of",
    "legislation of", "regulation of", "ordinance of", "decree of",
    "edict of", "proclamation of", "notification of", "directive of",
    "instruction of", "order of", "command of", "injunction of",
    "mandate of", "requirement of", "obligation of", "duty of",
    "responsibility of", "accountability of", "liability of", "culpability of",
    "blameworthiness of", "fault of", "negligence of", "recklessness of",
    "intent of", "purpose of", "motive of", "reason for",
    "cause of", "origin of", "source of", "basis of",
    "foundation of", "ground of", "justification for", "rationale for",
    "explanation for", "defense of", "excuse for", "pretext for",
    "alibi for", "vindication of", "exoneration of", "absolution of",
    "acquittal of", "discharge of", "release from", "liberation from",
    "emancipation from", "freedom from", "exemption from", "exception to",
    "derogation from", "deviation from", "departure from", "variation from",
    "modification of", "alteration of", "change to", "adjustment to",
    "adaptation of", "transformation of", "conversion of", "translation of",
    "interpretation of", "construction of", "reading of", "understanding of",
    "comprehension of", "appreciation of", "recognition of", "acknowledgment of",
    "admission of", "concession of", "confession of", "disclosure of",
    "revelation of", "exposure of", "discovery of", "finding of",
    "determination of", "resolution of", "decision on", "judgment on",
    "ruling on", "verdict on", "sentence on", "order on",
    "decree on", "pronouncement on", "declaration on", "announcement on",
    "publication of", "issuance of", "delivery of", "service of",
    "filing of", "submission of", "presentation of", "tender of",
    "offer of", "proposal of", "suggestion of", "recommendation of",
    "advice on", "counsel on", "guidance on", "direction on",
    "instruction on", "command on", "order on", "injunction on",
    "mandate on", "requirement on", "demand for", "request for",
    "petition for", "application for", "appeal for", "plea for",
    "prayer for", "suit for", "action for", "case for",
    "matter of", "issue of", "question of", "point of",
    "aspect of", "element of", "factor in", "component of",
    "ingredient of", "feature of", "characteristic of", "attribute of",
    "quality of", "property of", "trait of", "mark of",
    "sign of", "indication of", "evidence of", "proof of",
    "verification of", "confirmation of", "validation of", "authentication of",
    "certification of", "attestation of", "witnessing of", "observation of",
    "perception of", "view of", "opinion of", "belief of",
    "conviction of", "position of", "stance of", "attitude toward",
    "approach to", "method for", "technique for", "procedure for",
    "process for", "system for", "framework for", "model for",
    "paradigm for", "template for", "prototype for", "example of",
    "instance of", "case of", "illustration of", "demonstration of",
    "exposition of", "explanation of", "clarification of", "elucidation of",
    "simplification of", "breakdown of", "deconstruction of", "reconstruction of",
    "synthesis of", "integration of", "unification of", "harmonization of",
    "reconciliation of", "alignment of", "coordination of", "orchestration of",
    "management of", "administration of", "governance of", "regulation of",
    "control of", "supervision of", "oversight of", "monitoring of",
    "evaluation of", "assessment of", "appraisal of", "review of",
    "audit of", "inspection of", "scrutiny of", "examination of",
    "verification of", "validation of", "authentication of", "certification of",
    "accreditation of", "licensing of", "authorization of", "approval of",
    "sanction of", "endorsement of", "ratification of", "confirmation of",
    "affirmation of", "declaration of", "pronouncement of", "announcement of",
    "publication of", "dissemination of", "distribution of", "circulation of",
    "promulgation of", "enactment of", "legislation of", "regulation of",
    "ordinance of", "decree of", "edict of", "proclamation of",
    "notification of", "directive of", "instruction of", "order of",
    "command of", "injunction of", "mandate of", "requirement of",
    "obligation of", "duty of", "responsibility of", "accountability of",
    "liability of", "culpability of", "blameworthiness of", "fault of",
    "negligence of", "recklessness of", "intent of", "purpose of",
    "motive of", "reason for", "cause of", "origin of",
    "source of", "basis of", "foundation of", "ground of",
    "justification for", "rationale for", "explanation for", "defense of",
    "excuse for", "pretext for", "alibi for", "vindication of",
    "exoneration of", "absolution of", "acquittal of", "discharge of",
    "release from", "liberation from", "emancipation from", "freedom from",
    "exemption from", "exception to", "derogation from", "deviation from",
    "departure from", "variation from", "modification of", "alteration of",
    "change to", "adjustment to", "adaptation of", "transformation of",
    "conversion of", "translation of", "interpretation of", "construction of",
    "reading of", "understanding of", "comprehension of", "appreciation of",
    "recognition of", "acknowledgment of", "admission of", "concession of",
    "confession of", "disclosure of", "revelation of", "exposure of",
    "discovery of", "finding of", "determination of", "resolution of",
    "decision on", "judgment on", "ruling on", "verdict on",
    "sentence on", "order on", "decree on", "pronouncement on",
    "declaration on", "announcement on", "publication of", "issuance of",
    "delivery of", "service of", "filing of", "submission of",
    "presentation of", "tender of", "offer of", "proposal of",
    "suggestion of", "recommendation of", "advice on", "counsel on",
    "guidance on", "direction on", "instruction on", "command on",
    "order on", "injunction on", "mandate on", "requirement on",
    "demand for", "request for", "petition for", "application for",
    "appeal for", "plea for", "prayer for", "suit for",
    "action for", "case for", "matter of", "issue of",
    "question of", "point of", "aspect of", "element of",
    "factor in", "component of", "ingredient of", "feature of",
    "characteristic of", "attribute of", "quality of", "property of",
    "trait of", "mark of", "sign of", "indication of",
    "evidence of", "proof of", "verification of", "confirmation of",
    "validation of", "authentication of", "certification of", "attestation of",
    "witnessing of", "observation of", "perception of", "view of",
    "opinion of", "belief of", "conviction of", "position of",
    "stance of", "attitude toward", "approach to", "method for",
    "technique for", "procedure for", "process for", "system for",
    "framework for", "model for", "paradigm for", "template for",
    "prototype for", "example of", "instance of", "case of",
    "illustration of", "demonstration of", "exposition of", "explanation of",
    "clarification of", "elucidation of", "simplification of", "breakdown of",
    "deconstruction of", "reconstruction of", "synthesis of", "integration of",
    "unification of", "harmonization of", "reconciliation of", "alignment of",
    "coordination of", "orchestration of", "management of", "administration of",
    "governance of", "regulation of", "control of", "supervision of",
    "oversight of", "monitoring of", "evaluation of", "assessment of",
    "appraisal of", "review of", "audit of", "inspection of",
    "scrutiny of", "examination of", "verification of", "validation of",
    "authentication of", "certification of", "accreditation of", "licensing of",
    "authorization of", "approval of", "sanction of", "endorsement of",
    "ratification of", "confirmation of", "affirmation of", "declaration of",
    "pronouncement of", "announcement of", "publication of", "dissemination of",
    "distribution of", "circulation of", "promulgation of", "enactment of",
    "legislation of", "regulation of", "ordinance of", "decree of",
    "edict of", "proclamation of", "notification of", "directive of",
    "instruction of", "order of", "command of", "injunction of",
    "mandate of", "requirement of", "obligation of", "duty of",
    "responsibility of", "accountability of", "liability of", "culpability of",
    "blameworthiness of", "fault of", "negligence of", "recklessness of",
    "intent of", "purpose of", "motive of", "reason for",
    "cause of", "origin of", "source of", "basis of",
    "foundation of", "ground of", "justification for", "rationale for",
    "explanation for", "defense of", "excuse for", "pretext for",
    "alibi for", "vindication of", "exoneration of", "absolution of",
    "acquittal of", "discharge of", "release from", "liberation from",
    "emancipation from", "freedom from", "exemption from", "exception to",
    "derogation from", "deviation from", "departure from", "variation from",
    "modification of", "alteration of", "change to", "adjustment to",
    "adaptation of", "transformation of", "conversion of", "translation of",
    "interpretation of", "construction of", "reading of", "understanding of",
    "comprehension of", "appreciation of", "recognition of", "acknowledgment of",
    "admission of", "concession of", "confession of", "disclosure of",
    "revelation of", "exposure of", "discovery of", "finding of",
    "determination of", "resolution of", "decision on", "judgment on",
    "ruling on", "verdict on", "sentence on", "order on",
    "decree on", "pronouncement on", "declaration on", "announcement on",
    "publication of", "issuance of", "delivery of", "service of",
    "filing of", "submission of", "presentation of", "tender of",
    "offer of", "proposal of", "suggestion of", "recommendation of",
    "advice on", "counsel on", "guidance on", "direction on",
    "instruction on", "command on", "order on", "injunction on",
    "mandate on", "requirement on", "demand for", "request for",
    "petition for", "application for", "appeal for", "plea for",
    "prayer for", "suit for", "action for", "case for",
    "matter of", "issue of", "question of", "point of",
    "aspect of", "element of", "factor in", "component of",
    "ingredient of", "feature of", "characteristic of", "attribute of",
    "quality of", "property of", "trait of", "mark of",
    "sign of", "indication of", "evidence of", "proof of",
    "verification of", "confirmation of", "validation of", "authentication of",
    "certification of", "attestation of", "witnessing of", "observation of",
    "perception of", "view of", "opinion of", "belief of",
    "conviction of", "position of", "stance of", "attitude toward",
    "approach to", "method for", "technique for", "procedure for",
    "process for", "system for", "framework for", "model for",
    "paradigm for", "template for", "prototype for", "example of",
    "instance of", "case of", "illustration of", "demonstration of",
    "exposition of", "explanation of", "clarification of", "elucidation of",
    "simplification of", "breakdown of", "deconstruction of", "reconstruction of",
    "synthesis of", "integration of", "unification of", "harmonization of",
    "reconciliation of", "alignment of", "coordination of", "orchestration of",
    "management of", "administration of", "governance of", "regulation of",
    "control of", "supervision of", "oversight of", "monitoring of",
    "evaluation of", "assessment of", "appraisal of", "review of",
    "audit of", "inspection of", "scrutiny of", "examination of",
    "verification of", "validation of", "authentication of", "certification of",
    "accreditation of", "licensing of", "authorization of", "approval of",
    "sanction of", "endorsement of", "ratification of", "confirmation of",
    "affirmation of", "declaration of", "pronouncement of", "announcement of",
    "publication of", "dissemination of", "distribution of", "circulation of",
    "promulgation of", "enactment of", "legislation of", "regulation of",
    "ordinance of", "decree of", "edict of", "proclamation of",
    "notification of", "directive of", "instruction of", "order of",
    "command of", "injunction of", "mandate of", "requirement of",
    "obligation of", "duty of", "responsibility of", "accountability of",
    "liability of", "culpability of", "blameworthiness of", "fault of",
    "negligence of", "recklessness of", "intent of", "purpose of",
    "motive of", "reason for", "cause of", "origin of",
    "source of", "basis of", "foundation of", "ground of",
    "justification for", "rationale for", "explanation for", "defense of",
    "excuse for", "pretext for", "alibi for", "vindication of",
    "exoneration of", "absolution of", "acquittal of", "discharge of",
    "release from", "liberation from", "emancipation from", "freedom from",
    "exemption from", "exception to", "derogation from", "deviation from",
    "departure from", "variation from", "modification of", "alteration of",
    "change to", "adjustment to", "adaptation of", "transformation of",
    "conversion of", "translation of", "interpretation of", "construction of",
    "reading of", "understanding of", "comprehension of", "appreciation of",
    "recognition of", "acknowledgment of", "admission of", "concession of",
    "confession of", "disclosure of", "revelation of", "exposure of",
    "discovery of", "finding of", "determination of", "resolution of",
    "decision on", "judgment on", "ruling on", "verdict on",
    "sentence on", "order on", "decree on", "pronouncement on",
    "declaration on", "announcement on", "publication of", "issuance of",
    "delivery of", "service of", "filing of", "submission of",
    "presentation of", "tender of", "offer of", "proposal of",
    "suggestion of", "recommendation of", "advice on", "counsel on",
    "guidance on", "direction on", "instruction on", "command on",
    "order on", "injunction on", "mandate on", "requirement on",
    "demand for", "request for", "petition for", "application for",
    "appeal for", "plea for", "prayer for", "suit for",
    "action for", "case for", "matter of", "issue of",
    "question of", "point of", "aspect of", "element of",
    "factor in", "component of", "ingredient of", "feature of",
    "characteristic of", "attribute of", "quality of", "property of",
    "trait of", "mark of", "sign of", "indication of",
    "evidence of", "proof of", "verification of", "confirmation of",
    "validation of", "authentication of", "certification of", "attestation of",
    "witnessing of", "observation of", "perception of", "view of",
    "opinion of", "belief of", "conviction of", "position of",
    "stance of", "attitude toward", "approach to", "method for",
    "technique for", "procedure for", "process for", "system for",
    "framework for", "model for", "paradigm for", "template for",
    "prototype for", "example of", "instance of", "case of",
    "illustration of", "demonstration of", "exposition of", "explanation of",
    "clarification of", "elucidation of", "simplification of", "breakdown of",
    "deconstruction of", "reconstruction of", "synthesis of", "integration of",
    "unification of", "harmonization of", "reconciliation of", "alignment of",
    "coordination of", "orchestration of", "management of", "administration of",
    "governance of", "regulation of", "control of", "supervision of",
    "oversight of", "monitoring of", "evaluation of", "assessment of",
    "appraisal of", "review of", "audit of", "inspection of",
    "scrutiny of", "examination of", "verification of", "validation of",
    "authentication of", "certification of", "accreditation of", "licensing of",
    "authorization of", "approval of", "sanction of", "endorsement of",
    "ratification of", "confirmation of", "affirmation of", "declaration of",
    "pronouncement of", "announcement of", "publication of", "dissemination of",
    "distribution of", "circulation of", "promulgation of", "enactment of",
    "legislation of", "regulation of", "ordinance of", "decree of",
    "edict of", "proclamation of", "notification of", "directive of",
    "instruction of", "order of", "command of", "injunction of",
    "mandate of", "requirement of", "obligation of", "duty of",
    "responsibility of", "accountability of", "liability of", "culpability of",
    "blameworthiness of", "fault of", "negligence of", "recklessness of",
    "intent of", "purpose of", "motive of", "reason for",
    "cause of", "origin of", "source of", "basis of",
    "foundation of", "ground of", "justification for", "rationale for",
    "explanation for", "defense of", "excuse for", "pretext for",
    "alibi for", "vindication of", "exoneration of", "absolution of",
    "acquittal of", "discharge of", "release from", "liberation from",
    "emancipation from", "freedom from", "exemption from", "exception to",
    "derogation from", "deviation from", "departure from", "variation from",
    "modification of", "alteration of", "change to", "adjustment to",
    "adaptation of", "transformation of", "conversion of", "translation of",
    "interpretation of", "construction of", "reading of", "understanding of",
    "comprehension of", "appreciation of", "recognition of", "acknowledgment of",
    "admission of", "concession of", "confession of", "disclosure of",
    "revelation of", "exposure of", "discovery of", "finding of",
    "determination of", "resolution of", "decision on", "judgment on",
    "ruling on", "verdict on", "sentence on", "order on",
    "decree on", "pronouncement on", "declaration on", "announcement on",
    "publication of", "issuance of", "delivery of", "service of",
    "filing of", "submission of", "presentation of", "tender of",
    "offer of", "proposal of", "suggestion of", "recommendation of",
    "advice on", "counsel on", "guidance on", "direction on",
    "instruction on", "command on", "order on", "injunction on",
    "mandate on", "requirement on", "demand for", "request for",
    "petition for", "application for", "appeal for", "plea for",
    "prayer for", "suit for", "action for", "case for",
    "matter of", "issue of", "question of", "point of",
    "aspect of", "element of", "factor in", "component of",
    "ingredient of", "feature of", "characteristic of", "attribute of",
    "quality of", "property of", "trait of", "mark of",
    "sign of", "indication of", "evidence of", "proof of",
    "verification of", "confirmation of", "validation of", "authentication of",
    "certification of", "attestation of", "witnessing of", "observation of",
    "perception of", "view of", "opinion of", "belief of",
    "conviction of", "position of", "stance of", "attitude toward",
    "approach to", "method for", "technique for", "procedure for",
    "process for", "system for", "framework for", "model for","industrial","act","payment","law","court","bonus","order","government"
]
IRRELEVANT_RESPONSES = [
    "This question appears unrelated to the legal judgment document.",
    "The system only answers questions specifically about the uploaded court judgment.",
    "Your question doesn't appear relevant to this legal document.",
    "For questions about this specific judgment, please reference the case details.",
    "I can only answer questions about the legal judgment document."
]

# Initialize the summarization model
summarizer = None
tokenizer = None
model = None

try:
    logger.info("⏳ Loading summarization model...")
    device = -1  # Always use CPU
    logger.info("Using device: CPU")
    
    # tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
    tokenizer =MODEL_NAME
    # model = T5ForConditionalGeneration.from_pretrained(MODEL_NAME)
    model = MODEL_NAME
    summarizer = pipeline(
        "summarization",
        model=model,
        tokenizer=tokenizer,
        device=device
    )
    logger.info("✅ Model loaded successfully!")
except Exception as e:
    logger.error(f"❌ Failed to load model: {str(e)}\n{traceback.format_exc()}")
    raise

# Initialize QA vector stores cache
vectorstore_cache = {}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def log_memory_usage():
    logger.info("Memory usage logging disabled for CPU-only mode")

def chunk_text(text, chunk_size=CHUNK_SIZE):
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0

    for word in words:
        if current_length + len(word) < chunk_size:
            current_chunk.append(word)
            current_length += len(word)
        else:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_length = len(word)

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks

def summarize_chunk(chunk, max_length=DEFAULT_MAX_LENGTH, min_length=DEFAULT_MIN_LENGTH):
    try:
        return summarizer(
            chunk,
            max_length=max_length,
            min_length=min_length,
            length_penalty=1.5,
            num_beams=4,
            no_repeat_ngram_size=3,
            early_stopping=True
        )[0]['summary_text']
    except Exception as e:
        logger.error(f"Error summarizing chunk: {str(e)}")
        return ""

def parallel_summarize(text, max_length=None, min_length=None):
    if not text.strip():
        return ""
        
    word_count = len(text.split())
    if max_length is None:
        max_length = min(DEFAULT_MAX_LENGTH + (word_count // 100), 512)
    if min_length is None:
        min_length = min(DEFAULT_MIN_LENGTH + (word_count // 200), 256)
    
    chunks = chunk_text(text)
    
    with ThreadPoolExecutor() as executor:
        summaries = list(executor.map(
            lambda c: summarize_chunk(c, max_length, min_length),
            chunks
        ))
    
    final_summary = " ".join([s for s in summaries if s])
    final_summary = re.sub(r'\s+([.,;:])', r'\1', final_summary)
    final_summary = re.sub(r'\.\s+\.', '.', final_summary)
    final_summary = re.sub(r'\s+', ' ', final_summary).strip()
    
    return final_summary

def preprocess_text(text):
    try:
        if not text or not isinstance(text, str):
            return ""
            
        text = re.sub(r'Page\s*\d+\s*of\s*\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\n\d+\n', '\n', text)
        
        noise_patterns = [
            r'This\s+document\s+was\s+signed\s+electronically.*',
            r'Electronic\s+signature.*',
            r'IN\s+THE\s+(?:SUPREME\s+)?COURT\s+OF\s+.*\n',
            r'CASE\s+NO[.:].*\n',
            r'BEFORE[.:].*\n',
            r'JUDGMENT\s+RESERVED\s+ON[.:].*\n',
            r'PRESENT[.:].*\n'
        ]
        
        for pattern in noise_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if (re.fullmatch(r'\d+', line) or 
                len(line) < 25 or 
                line.startswith(('§', '©', 'Page', 'http'))):
                continue
                
            cleaned_lines.append(line)
        
        return '\n\n'.join(cleaned_lines) or ""
    except Exception as e:
        logger.error(f"Error in preprocessing: {str(e)}\n{traceback.format_exc()}")
        return text if isinstance(text, str) else ""

def get_processed_filename(original_filename):
    base = os.path.splitext(os.path.basename(original_filename))[0]
    return f"processed_{secure_filename(base)}.txt"

def get_processed_path(filename):
    processed_filename = get_processed_filename(filename)
    return os.path.normpath(os.path.join(app.config['PROCESSED_FOLDER'], processed_filename))

def extract_text_from_file(filepath, filename):
    try:
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == '.pdf':
            text = ""
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
            return text.strip()
        
        elif ext in ['.doc', '.docx']:
            doc = docx.Document(filepath)
            return "\n".join(para.text for para in doc.paragraphs if para.text)
        
        elif ext == '.txt':
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        
        raise ValueError(f"Unsupported file type: {ext}")
    
    except Exception as e:
        logger.error(f"Extraction failed for {filename}: {str(e)}")
        raise

def create_vector_store(text, filename):
    try:
        processed_path = get_processed_path(filename)
        os.makedirs(os.path.dirname(processed_path), exist_ok=True)
        
        with open(processed_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        if not os.path.exists(processed_path):
            raise IOError(f"Failed to create processed file at {processed_path}")
        
        with open(processed_path, 'r', encoding='utf-8') as f:
            text = f.read()
        documents = [Document(page_content=text, metadata={"source": processed_path})]
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        chunks = text_splitter.split_documents(documents)
        
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-mpnet-base-v2",
            model_kwargs={'device': 'cpu'}
        )
        
        vectorstore = FAISS.from_documents(chunks, embeddings)
        return vectorstore, None
            
    except Exception as e:
        logger.error(f"Error in create_vector_store: {str(e)}")
        return None, str(e)

def contains_legal_terms(text):
    if not text:
        return False
    text_lower = text.lower()
    return sum(term in text_lower for term in LEGAL_TERMS) >= 2

def is_relevant_response(query, docs_and_scores):
    if not docs_and_scores or not any(score >= 0.8 for _, score in docs_and_scores):
        return False

    query_lower = query.lower()
    if not (query_lower.endswith('?') or 
            any(word in query_lower for word in ['who', 'what', 'when', 'where', 'why', 'how', 'explain'])):
        return False

    for doc, _ in docs_and_scores:
        if contains_legal_terms(doc.page_content) and contains_legal_terms(query):
            return True
    return False

def format_answer(doc, score):
    content = doc.page_content
    return {
        "content": content,
        "score":float(score)
}

@app.route('/health', methods=['GET'])
def health_check():
    log_memory_usage()
    return jsonify({
        "status": "healthy",
        "model_loaded": summarizer is not None,
        "device": "cpu",
        "chunk_size": CHUNK_SIZE,
        "default_max_length": DEFAULT_MAX_LENGTH,
        "default_min_length": DEFAULT_MIN_LENGTH
    })

@app.route('/summarize', methods=['POST'])
def summarize():
    if summarizer is None:
        return jsonify({
            "error": "Model not loaded",
            "summary": "",
            "status": "error"
        }), 503
    
    if 'file' not in request.files:
        return jsonify({
            "error": "No file uploaded",
            "summary": "",
            "status": "error"
        }), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({
            "error": "No selected file",
            "summary": "",
            "status": "error"
        }), 400
    
    if not file or not allowed_file(file.filename):
        return jsonify({
            "error": "Invalid file type. Allowed: pdf, doc, docx, txt",
            "summary": "",
            "status": "error"
        }), 400
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    preprocessed_path = os.path.join(app.config['PREPROCESSED_FOLDER'], f"preprocessed_{filename}.txt")
    
    try:
        file.save(filepath)
        logger.info(f"Processing file: {filename}")
        log_memory_usage()
        
        raw_text = extract_text_from_file(filepath, filename)
        if not raw_text or not raw_text.strip():
            logger.error(f"Empty text extracted from {filename}")
            return jsonify({
                "error": "Empty file or could not extract text",
                "summary": "",
                "status": "error"
            }), 400
            
        cleaned_text = preprocess_text(raw_text)
        logger.info(f"Text length: {len(cleaned_text)} chars, {len(cleaned_text.split())} words")
        
        with open(preprocessed_path, 'w', encoding='utf-8') as f:
            f.write(cleaned_text)
        
        start_time = time.time()
        summary = parallel_summarize(cleaned_text)
        processing_time = time.time() - start_time
        
        if not summary:
            raise ValueError("Failed to generate summary - empty result")
        
        logger.info(f"Generated summary in {processing_time:.2f} seconds")
        logger.info(f"Summary length: {len(summary.split())} words")
        
        return jsonify({
            "summary": summary,
            "filename": filename,
            "processing_time": f"{processing_time:.2f} seconds",
            "word_count": len(cleaned_text.split()),
            "summary_length": len(summary.split()),
            "status": "success"
        })
        
    except Exception as e:
        logger.error(f"Error processing {filename}: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            "error": "Processing failed",
            "details": str(e),
            "filename": filename,
            "summary": "",
            "status": "error"
        }), 500
    finally:
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
            except Exception as e:
                logger.error(f"Error removing file {filepath}: {str(e)}")
        log_memory_usage()

@app.route('/upload', methods=['POST'])
def upload_document():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded", "status": "error"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file", "status": "error"}), 400
    
    if not allowed_file(file.filename):
        return jsonify({
            "error": "Invalid file type. Allowed: pdf, doc, docx, txt",
            "status": "error"
        }), 400
    
    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        file.save(filepath)
        
        raw_text = extract_text_from_file(filepath, filename)
        if not raw_text:
            return jsonify({
                "error": "Could not extract text from document",
                "status": "error"
            }), 400
            
        cleaned_text = preprocess_text(raw_text)
        vectorstore, error = create_vector_store(cleaned_text, filename)
        if error:
            return jsonify({"error": error, "status": "error"}), 500
        
        vectorstore_path = os.path.join(app.config['PROCESSED_FOLDER'], f"vectorstore_{filename}")
        vectorstore.save_local(vectorstore_path)
        vectorstore_cache[filename] = vectorstore
        
        return jsonify({
            "message": "Document processed successfully",
            "filename": filename,
            "status": "success"
        })
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            "error": "Document processing failed",
            "details": str(e),
            "status": "error"
        }), 500
    finally:
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
            except Exception as e:
                logger.error(f"Error removing file {filepath}: {str(e)}")

@app.route('/ask', methods=['POST'])
def ask_question():
    try:
        if 'file' not in request.files:
            return jsonify({
                "error": "Missing file",
                "status": "error"
            }), 400

        file = request.files['file']
        question = request.form.get('question', '').strip()

        if not file or file.filename == '':
            return jsonify({"error": "No file selected", "status": "error"}), 400

        if len(question.split()) < 3:
            return jsonify({
                "error": "Please ask a more detailed question (minimum 3 words)",
                "status": "success"
            }), 200

        filename = secure_filename(file.filename)
        vectorstore = vectorstore_cache.get(filename)
        
        if vectorstore is None:
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            raw_text = extract_text_from_file(filepath, filename)
            if not raw_text:
                return jsonify({
                    "error": "Could not extract text from document",
                    "status": "error"
                }), 400

            cleaned_text = preprocess_text(raw_text)
            vectorstore, error = create_vector_store(cleaned_text, filename)
            if error:
                return jsonify({"error": error, "status": "error"}), 500
            
            if os.path.exists(filepath):
                os.remove(filepath)

        docs_and_scores = vectorstore.similarity_search_with_score(question, k=5)
        
        if not is_relevant_response(question, docs_and_scores):
            return jsonify({
                "answer": random.choice(IRRELEVANT_RESPONSES),
                "sections": [],
                "isRelevant": False,
                "filename": filename,
                "status": "success"
            })
        
        relevant_sections = []
        for doc, score in docs_and_scores:
            if score >= 0.8 and contains_legal_terms(doc.page_content):
                relevant_sections.append(format_answer(doc, score))

        return jsonify({
            "answer": "Here are the relevant sections from the document:",
            "sections": relevant_sections,
            "filename": filename,
            "isRelevant": True,
            "status": "success"
        })

    except Exception as e:
        logger.error(f"Error answering question: {str(e)}\n{traceback.format_exc()}")
        return jsonify({
            "error": "Failed to process question",
            "details": str(e),
            "status": "error"
        }), 500
app.config.from_mapping(
    MYMEMORY_URL='https://api.mymemory.translated.net/get',
    LIBRE_URL='https://libretranslate.de/translate',
    RATE_LIMIT=10,  # requests per minute
    CACHE_SIZE=1000,
    REQUEST_TIMEOUT=10  # seconds
)

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Rate limiting storage
request_timestamps = {}

def rate_limited(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        ip = request.remote_addr
        now = datetime.now()
        
        # Clean up old timestamps
        request_timestamps[ip] = [
            ts for ts in request_timestamps.get(ip, []) 
            if now - ts < timedelta(minutes=1)
        ]
        
        # Check rate limit
        if len(request_timestamps.get(ip, [])) >= app.config['RATE_LIMIT']:
            logger.warning(f"Rate limit exceeded for IP: {ip}")
            return jsonify({
                'error': 'Rate limit exceeded',
                'message': f'Please wait and try again. Limit is {app.config["RATE_LIMIT"]} requests per minute.'
            }), 429
        
        # Add new timestamp
        request_timestamps.setdefault(ip, []).append(now)
        
        return f(*args, **kwargs)
    return decorated_function

@lru_cache(maxsize=app.config['CACHE_SIZE'])
def translate_text(text, target_lang):
    """
    Translates text from English to target language using MyMemory or LibreTranslate as fallback.
    
    Args:
        text (str): Text to translate (must be in English)
        target_lang (str): Target language code (e.g., 'es', 'fr', 'de')
    
    Returns:
        str: Translated text
    
    Raises:
        Exception: If translation fails
    """
    # Try MyMemory first
    try:
        response = requests.get(
            app.config['MYMEMORY_URL'],
            params={'q': text, 'langpair': f'en|{target_lang}'},
            timeout=app.config['REQUEST_TIMEOUT']
        )
        if response.status_code == 200:
            data = response.json()
            if 'responseData' in data and 'translatedText' in data['responseData']:
                return data['responseData']['translatedText']
    except requests.exceptions.RequestException as e:
        logger.warning(f"MyMemory translation failed: {str(e)}")
    
    # Fallback to LibreTranslate
    try:
        response = requests.post(
            app.config['LIBRE_URL'],
            json={'q': text, 'source': 'en', 'target': target_lang},
            timeout=app.config['REQUEST_TIMEOUT']
        )
        if response.status_code == 200:
            data = response.json()
            if 'translatedText' in data:
                return data['translatedText']
        elif response.status_code == 429:
            raise Exception('LibreTranslate rate limit exceeded')
    except requests.exceptions.RequestException as e:
        logger.warning(f"LibreTranslate translation failed: {str(e)}")
        raise Exception('All translation services failed') from e
    
    raise Exception('Translation failed - no valid response from services')

@app.errorhandler(Exception)
def handle_exception(e):
    # Pass through HTTP errors
    if isinstance(e, HTTPException):
        return e
    
    # Log the error
    logger.error(f"Unexpected error: {str(e)}", exc_info=True)
    
    # Return JSON response
    return jsonify({
        'error': 'Internal server error',
        'message': str(e)
    }), 500


@app.route('/translate', methods=['POST'])
@rate_limited
def translate_endpoint():
    """
    Translation endpoint
    
    Expected JSON payload:
    {
        "text": "text to translate",
        "lang": "target_language_code",
        "chunked": false (optional)  # whether to split long texts
    }
    """
    MAX_LENGTH = 500  # Define your maximum length for single translation
    
    # Validate input
    if not request.is_json:
        return jsonify({'error': 'Request must be JSON'}), 400
    
    data = request.get_json()
    text = data.get('text')
    lang = data.get('lang')
    chunked = data.get('chunked', False)  # Default to False if not provided
    
    if not text or not lang:
        return jsonify({
            'error': 'Missing required parameters',
            'message': 'Both "text" and "lang" parameters are required'
        }), 400
    
    if not isinstance(text, str) or not isinstance(lang, str):
        return jsonify({
            'error': 'Invalid parameters',
            'message': 'Both "text" and "lang" must be strings'
        }), 400
    
    # SOLUTION 1: STRICT LENGTH CHECK (uncomment to use)
    # if len(text) > MAX_LENGTH:
    #     return jsonify({
    #         'error': 'Text too long',
    #         'message': f'Maximum allowed length is {MAX_LENGTH} characters',
    #         'submitted_length': len(text),
    #         'suggestion': 'Try splitting your text into smaller chunks'
    #     }), 400

    # SOLUTION 2: CHUNKED TRANSLATION (default implementation)
    def translate_in_chunks(text, lang, chunk_size=MAX_LENGTH):
        chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
        translated_chunks = []
        for chunk in chunks:
            translated_chunks.append(translate_text(chunk, lang))
        return ' '.join(translated_chunks)

    # Perform translation
    try:
        start_time = time.time()
        
        if chunked or len(text) > MAX_LENGTH:
            translated_text = translate_in_chunks(text, lang)
            was_chunked = True
        else:
            translated_text = translate_text(text, lang)
            was_chunked = False
            
        duration = time.time() - start_time
        
        logger.info(f"Translated text (length: {len(text)}) to {lang} in {duration:.2f}s" + 
                   (" (chunked)" if was_chunked else ""))
        
        return jsonify({
            'translation': translated_text,
            'original_text': text,
            'source_lang': 'en',  # assuming source is English
            'target_lang': lang,
            'duration_seconds': round(duration, 2),
            'was_chunked': was_chunked,
            'original_length': len(text)
        })
        
    except Exception as e:
        logger.error(f"Translation failed: {str(e)}")
        return jsonify({
            'error': 'Translation failed',
            'message': str(e),
            'original_text': text,
            'target_lang': lang,
            'original_length': len(text)
        }), 500
@app.route('/healthy', methods=['GET'])
def health_check_qa():
    return jsonify({
        "status": "healthy",
        "message": "Legal QA system is running"
    })

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)